"""DOCX renderer using python-docx.

Walks ProseMirror JSON directly and builds a Word document with proper
manuscript formatting (page breaks, headers/footers, template styles).
"""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from app.modules.export.renderers.base import BaseRenderer, RendererNode


_ALIGN_MAP: dict[str, WD_ALIGN_PARAGRAPH] = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DocxRenderer(BaseRenderer):
    def render(
        self,
        nodes: list[RendererNode],
        template_config: dict,
        project_title: str,
    ) -> bytes:
        doc = Document()
        self._apply_template(doc, template_config, project_title)

        first_content_node = True
        for node in nodes:
            if node.node_type == "folder":
                # Folders are structural only — add a section divider heading
                if not first_content_node:
                    doc.add_page_break()
                para = doc.add_heading(node.title, level=1)
                para.style.font.name = template_config.get("font_family", "Times New Roman")
                first_content_node = False
                continue

            if node.node_type == "chapter":
                if not first_content_node and template_config.get("chapter_break", True):
                    doc.add_page_break()
                heading = doc.add_heading(node.title, level=1)
                heading.style.font.name = template_config.get("font_family", "Times New Roman")
                first_content_node = False

            elif node.node_type == "scene":
                if not first_content_node:
                    sep = template_config.get("scene_separator", "#")
                    sep_para = doc.add_paragraph(sep)
                    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_content_node = False

            elif node.node_type in ("front_matter", "back_matter"):
                if not first_content_node:
                    doc.add_page_break()
                doc.add_heading(node.title, level=1)
                first_content_node = False

            if node.content:
                walker = _DocxWalker(doc, template_config)
                walker.walk(node.content)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def _apply_template(self, doc: Document, config: dict, project_title: str) -> None:
        """Apply margin, font, line-spacing, and header/footer settings."""
        font_family = config.get("font_family", "Times New Roman")
        font_size = config.get("font_size", 12)
        line_spacing = config.get("line_spacing", 2.0)
        margin_top = config.get("margin_top", 1.0)
        margin_bottom = config.get("margin_bottom", 1.0)
        margin_left = config.get("margin_left", 1.0)
        margin_right = config.get("margin_right", 1.0)

        section = doc.sections[0]
        section.top_margin = Inches(margin_top)
        section.bottom_margin = Inches(margin_bottom)
        section.left_margin = Inches(margin_left)
        section.right_margin = Inches(margin_right)

        # Default paragraph style
        style = doc.styles["Normal"]
        style.font.name = font_family
        style.font.size = Pt(font_size)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing

        # Header (project title)
        if config.get("header_text") and config.get("header_text") != "":
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = project_title
            header_para.style.font.name = font_family
            header_para.style.font.size = Pt(font_size)

        # Footer (page numbers)
        if config.get("page_numbers", True):
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._add_page_number(footer_para)

    @staticmethod
    def _add_page_number(paragraph: Paragraph) -> None:
        """Insert a PAGE field into a paragraph."""
        run = paragraph.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)  # noqa: SLF001
        run._r.append(instr_text)  # noqa: SLF001
        run._r.append(fld_char2)  # noqa: SLF001


class _DocxWalker:
    """Walks a ProseMirror JSON tree and appends content to a Document."""

    def __init__(self, doc: Document, config: dict) -> None:
        self._doc = doc
        self._font = config.get("font_family", "Times New Roman")
        self._size = Pt(config.get("font_size", 12))

    def walk(self, node: dict) -> None:
        node_type = node.get("type", "")
        content: list[dict] = node.get("content", [])
        attrs: dict = node.get("attrs", {})

        match node_type:
            case "doc":
                for child in content:
                    self.walk(child)

            case "paragraph":
                para = self._doc.add_paragraph()
                self._set_align(para, attrs)
                for child in content:
                    self._inline(child, para)

            case "heading":
                level = min(attrs.get("level", 1), 9)
                para = self._doc.add_heading(level=level)
                self._set_align(para, attrs)
                for child in content:
                    self._inline(child, para)

            case "bulletList" | "orderedList":
                for child in content:
                    self.walk(child)

            case "listItem":
                para = self._doc.add_paragraph(style="List Bullet")
                for child in content:
                    # listItem can contain paragraphs; flatten them
                    if child.get("type") == "paragraph":
                        for inline in child.get("content", []):
                            self._inline(inline, para)
                    else:
                        self.walk(child)

            case "blockquote":
                for child in content:
                    if child.get("type") == "paragraph":
                        para = self._doc.add_paragraph()
                        para.paragraph_format.left_indent = Inches(0.5)
                        para.paragraph_format.right_indent = Inches(0.5)
                        self._set_align(para, child.get("attrs", {}))
                        for inline in child.get("content", []):
                            self._inline(inline, para)
                    else:
                        self.walk(child)

            case "codeBlock":
                for child in content:
                    if child.get("type") == "text":
                        para = self._doc.add_paragraph()
                        run = para.add_run(child.get("text", ""))
                        run.font.name = "Courier New"
                        run.font.size = self._size

            case "horizontalRule":
                para = self._doc.add_paragraph()
                self._add_border(para)

            case "hardBreak":
                pass  # handled inline

            case "image":
                # Images in exported manuscripts are not common; add placeholder
                para = self._doc.add_paragraph("[Image]")
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            case "table":
                self._walk_table(node)

            case _:
                for child in content:
                    self.walk(child)

    def _inline(self, node: dict, para: Paragraph) -> None:
        node_type = node.get("type", "")
        if node_type == "text":
            text = node.get("text", "")
            marks: list[dict] = node.get("marks", [])
            run = para.add_run(text)
            run.font.name = self._font
            run.font.size = self._size
            self._apply_marks(run, marks)
        elif node_type == "hardBreak":
            run = para.add_run()
            run.add_break()
        # Other inline types (image, etc.) are skipped in inline context

    def _apply_marks(self, run: object, marks: list[dict]) -> None:
        for mark in marks:
            mark_type = mark.get("type", "")
            attrs: dict = mark.get("attrs", {})
            match mark_type:
                case "bold":
                    run.bold = True  # type: ignore[union-attr]
                case "italic":
                    run.italic = True  # type: ignore[union-attr]
                case "underline":
                    run.underline = True  # type: ignore[union-attr]
                case "strike":
                    run.font.strike = True  # type: ignore[union-attr]
                case "code":
                    run.font.name = "Courier New"  # type: ignore[union-attr]
                case "subscript":
                    run.font.subscript = True  # type: ignore[union-attr]
                case "superscript":
                    run.font.superscript = True  # type: ignore[union-attr]
                case "textStyle":
                    if color := attrs.get("color"):
                        from docx.shared import RGBColor
                        # color may be #rrggbb or rgb() — handle hex
                        if color.startswith("#") and len(color) == 7:
                            r = int(color[1:3], 16)
                            g = int(color[3:5], 16)
                            b = int(color[5:7], 16)
                            run.font.color.rgb = RGBColor(r, g, b)  # type: ignore[union-attr]
                    if font_size := attrs.get("fontSize"):
                        try:
                            size_pt = float(str(font_size).replace("pt", ""))
                            run.font.size = Pt(size_pt)  # type: ignore[union-attr]
                        except ValueError:
                            pass

    def _set_align(self, para: Paragraph, attrs: dict) -> None:
        align = attrs.get("textAlign", "left")
        para.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    def _add_border(self, para: Paragraph) -> None:
        """Add a bottom border to simulate a horizontal rule."""
        pPr = para._p.get_or_add_pPr()  # noqa: SLF001
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _walk_table(self, table_node: dict) -> None:
        rows = [r for r in table_node.get("content", []) if r.get("type") == "tableRow"]
        if not rows:
            return
        cols = max(len(r.get("content", [])) for r in rows)
        table = self._doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for r_idx, row_node in enumerate(rows):
            cells = row_node.get("content", [])
            for c_idx, cell_node in enumerate(cells):
                if c_idx >= cols:
                    break
                cell = table.cell(r_idx, c_idx)
                cell.text = ""  # clear default empty paragraph
                para = cell.paragraphs[0]
                for inline in cell_node.get("content", []):
                    if inline.get("type") == "paragraph":
                        for child in inline.get("content", []):
                            self._inline(child, para)
                    else:
                        self._inline(inline, para)
