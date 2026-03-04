"""
Tests for the export module — renderers, job lifecycle, and API endpoints.

Renderer tests run without any external services.
API tests require a running database (integration).
"""

import io
import json
import uuid
import zipfile
from unittest.mock import MagicMock, patch

import pytest
from httpx import AsyncClient

from app.modules.export.renderers.base import RendererNode
from app.modules.export.renderers.docx_renderer import DocxRenderer
from app.modules.export.renderers.epub_renderer import EpubRenderer
from app.modules.export.renderers.pdf_renderer import PdfRenderer
from app.modules.export.renderers.utils import prosemirror_to_html

# ---------------------------------------------------------------------------
# Sample document
# ---------------------------------------------------------------------------

_SAMPLE_DOC = {
    "type": "doc",
    "content": [
        {
            "type": "heading",
            "attrs": {"level": 1},
            "content": [{"type": "text", "text": "Chapter One"}],
        },
        {
            "type": "paragraph",
            "content": [
                {"type": "text", "text": "Hello "},
                {
                    "type": "text",
                    "text": "bold world",
                    "marks": [{"type": "bold"}],
                },
                {"type": "text", "text": "."},
            ],
        },
        {
            "type": "paragraph",
            "content": [
                {
                    "type": "text",
                    "text": "Italic text.",
                    "marks": [{"type": "italic"}],
                }
            ],
        },
    ],
}

_TEMPLATE_CONFIG = {
    "font_family": "Times New Roman",
    "font_size": 12,
    "line_spacing": 2.0,
    "margin_top": 1.0,
    "margin_bottom": 1.0,
    "margin_left": 1.0,
    "margin_right": 1.0,
    "scene_separator": "#",
    "page_numbers": True,
    "chapter_break": True,
    "header_text": "Test Manuscript",
}

_NODES = [
    RendererNode(id="1", node_type="chapter", title="Chapter One", depth=0, content=_SAMPLE_DOC),
    RendererNode(id="2", node_type="scene", title="Scene One", depth=1, content=_SAMPLE_DOC),
]


# ---------------------------------------------------------------------------
# ProseMirror → HTML converter tests
# ---------------------------------------------------------------------------


class TestProsemirrorToHtml:
    def test_paragraph(self):
        doc = {"type": "doc", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Hello"}]}]}
        html = prosemirror_to_html(doc)
        assert "<p>" in html
        assert "Hello" in html

    def test_heading(self):
        node = {"type": "heading", "attrs": {"level": 2}, "content": [{"type": "text", "text": "Title"}]}
        html = prosemirror_to_html(node)
        assert html.startswith("<h2")
        assert "Title" in html

    def test_bold_mark(self):
        node = {"type": "text", "text": "bold", "marks": [{"type": "bold"}]}
        html = prosemirror_to_html(node)
        assert "<strong>bold</strong>" == html

    def test_italic_mark(self):
        node = {"type": "text", "text": "italic", "marks": [{"type": "italic"}]}
        html = prosemirror_to_html(node)
        assert "<em>italic</em>" == html

    def test_link_mark(self):
        node = {
            "type": "text",
            "text": "click",
            "marks": [{"type": "link", "attrs": {"href": "https://example.com", "target": "_blank"}}],
        }
        html = prosemirror_to_html(node)
        assert 'href="https://example.com"' in html
        assert "click" in html

    def test_hardbreak(self):
        node = {"type": "hardBreak"}
        assert prosemirror_to_html(node) == "<br>"

    def test_horizontal_rule(self):
        node = {"type": "horizontalRule"}
        assert prosemirror_to_html(node) == "<hr>"

    def test_bullet_list(self):
        node = {
            "type": "bulletList",
            "content": [
                {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "item"}]}]}
            ],
        }
        html = prosemirror_to_html(node)
        assert "<ul>" in html
        assert "<li>" in html
        assert "item" in html

    def test_html_escaping(self):
        node = {"type": "text", "text": "<script>alert('xss')</script>"}
        html = prosemirror_to_html(node)
        assert "<script>" not in html
        assert "&lt;script&gt;" in html

    def test_unknown_node_renders_children(self):
        node = {
            "type": "custom_block",
            "content": [{"type": "text", "text": "inner"}],
        }
        html = prosemirror_to_html(node)
        assert "inner" in html


# ---------------------------------------------------------------------------
# DOCX renderer tests
# ---------------------------------------------------------------------------


class TestDocxRenderer:
    def test_produces_bytes(self):
        renderer = DocxRenderer()
        output = renderer.render(_NODES, _TEMPLATE_CONFIG, "Test Manuscript")
        assert isinstance(output, bytes)
        assert len(output) > 100

    def test_valid_docx_structure(self):
        from docx import Document

        renderer = DocxRenderer()
        output = renderer.render(_NODES, _TEMPLATE_CONFIG, "Test Manuscript")
        doc = Document(io.BytesIO(output))
        # Should have paragraphs with some text
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Chapter One" in all_text or len(doc.paragraphs) > 0

    def test_empty_nodes(self):
        renderer = DocxRenderer()
        output = renderer.render([], _TEMPLATE_CONFIG, "Empty")
        assert isinstance(output, bytes)


# ---------------------------------------------------------------------------
# PDF renderer tests
# ---------------------------------------------------------------------------


class TestPdfRenderer:
    def test_produces_pdf_bytes(self):
        renderer = PdfRenderer()
        try:
            output = renderer.render(_NODES, _TEMPLATE_CONFIG, "Test Manuscript")
            assert isinstance(output, bytes)
            assert output[:4] == b"%PDF"
        except (ImportError, OSError):
            pytest.skip("WeasyPrint system libraries not available")

    def test_empty_nodes_produces_pdf(self):
        renderer = PdfRenderer()
        try:
            output = renderer.render([], _TEMPLATE_CONFIG, "Empty")
            assert output[:4] == b"%PDF"
        except (ImportError, OSError):
            pytest.skip("WeasyPrint system libraries not available")


# ---------------------------------------------------------------------------
# ePub renderer tests
# ---------------------------------------------------------------------------


class TestEpubRenderer:
    def test_produces_bytes(self):
        renderer = EpubRenderer()
        output = renderer.render(_NODES, _TEMPLATE_CONFIG, "Test Manuscript")
        assert isinstance(output, bytes)
        assert len(output) > 100

    def test_valid_epub_is_zip(self):
        renderer = EpubRenderer()
        output = renderer.render(_NODES, _TEMPLATE_CONFIG, "Test Manuscript")
        buf = io.BytesIO(output)
        assert zipfile.is_zipfile(buf)

    def test_epub_has_container_xml(self):
        renderer = EpubRenderer()
        output = renderer.render(_NODES, _TEMPLATE_CONFIG, "Test Manuscript")
        buf = io.BytesIO(output)
        with zipfile.ZipFile(buf) as z:
            names = z.namelist()
            assert "META-INF/container.xml" in names

    def test_epub_has_chapter_content(self):
        renderer = EpubRenderer()
        output = renderer.render(_NODES, _TEMPLATE_CONFIG, "Test Manuscript")
        buf = io.BytesIO(output)
        with zipfile.ZipFile(buf) as z:
            xhtml_files = [n for n in z.namelist() if n.endswith(".xhtml")]
            assert len(xhtml_files) > 0

    def test_empty_nodes(self):
        renderer = EpubRenderer()
        output = renderer.render([], _TEMPLATE_CONFIG, "Empty")
        assert isinstance(output, bytes)


# ---------------------------------------------------------------------------
# API endpoint tests (integration — require running DB)
# ---------------------------------------------------------------------------


@pytest.mark.integration
async def test_list_templates(async_client: AsyncClient, auth_headers: dict) -> None:
    response = await async_client.get("/api/v1/export/templates", headers=auth_headers)
    assert response.status_code == 200
    templates = response.json()
    assert isinstance(templates, list)
    assert len(templates) >= 3  # 3 seeded defaults
    formats = {t["format"] for t in templates}
    assert "docx" in formats
    assert "pdf" in formats
    assert "epub" in formats


@pytest.mark.integration
async def test_create_export_job(async_client: AsyncClient, auth_headers: dict) -> None:
    # Create a project first
    proj = (
        await async_client.post("/api/v1/projects", json={"title": "Export Test"}, headers=auth_headers)
    ).json()
    project_id = proj["id"]

    # Create export job (mocking the Celery task)
    with patch("app.tasks.export_tasks.export_document") as mock_task:
        mock_task.delay = MagicMock()
        response = await async_client.post(
            f"/api/v1/projects/{project_id}/export",
            json={"format": "docx", "scope": {"type": "full"}},
            headers=auth_headers,
        )

    assert response.status_code == 202
    job = response.json()
    assert job["status"] == "pending"
    assert job["format"] == "docx"
    assert job["project_id"] == project_id


@pytest.mark.integration
async def test_list_export_jobs(async_client: AsyncClient, auth_headers: dict) -> None:
    proj = (
        await async_client.post("/api/v1/projects", json={"title": "Export List Test"}, headers=auth_headers)
    ).json()
    project_id = proj["id"]

    response = await async_client.get(f"/api/v1/projects/{project_id}/export", headers=auth_headers)
    assert response.status_code == 200
    assert isinstance(response.json(), list)


@pytest.mark.integration
async def test_invalid_format_rejected(async_client: AsyncClient, auth_headers: dict) -> None:
    proj = (
        await async_client.post("/api/v1/projects", json={"title": "Bad Format"}, headers=auth_headers)
    ).json()
    response = await async_client.post(
        f"/api/v1/projects/{proj['id']}/export",
        json={"format": "rtf", "scope": {"type": "full"}},
        headers=auth_headers,
    )
    assert response.status_code == 422
